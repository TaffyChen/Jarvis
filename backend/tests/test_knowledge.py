from __future__ import annotations

import pytest

from app.services.knowledge import (
    delete_kb_document,
    get_kb_document,
    list_kb_documents,
    preview_kb_chunks,
    save_kb_document,
    upload_kb_document,
)
from app.infrastructure.kb.extract import extract_markdown, md_path_for_upload
from app.core.config import settings
from app.infrastructure.kb.chunk import chunk_markdown
from app.infrastructure.kb.search import keyword_rank, rrf_fuse
from app.infrastructure.kb.index import collect_chunks, reset_store, write_kb_meta


@pytest.fixture()
def isolated_kb(tmp_path, isolated_data_dir):
    old_knowledge = settings.knowledge_dir
    try:
        object.__setattr__(settings, "knowledge_dir", tmp_path / "knowledge")
        settings.knowledge_dir.mkdir(parents=True, exist_ok=True)
        reset_store()
        yield tmp_path
    finally:
        reset_store()
        object.__setattr__(settings, "knowledge_dir", old_knowledge)


def test_chunk_markdown_keeps_heading_path():
    text = "# 总则\n\n先看市场情绪。\n\n## 五灯\n\n仓位按五灯加减。第二句补充细节。" * 8
    chunks = chunk_markdown(text, "demo.md", chunk_size=80, overlap=20, min_len=8)
    assert chunks
    assert any("五灯" in (c.meta.get("heading") or "") for c in chunks)
    assert all(c.source == "demo.md" for c in chunks)


def test_rrf_prefers_overlap():
    vector = [
        {"id": "a", "text": "向量命中A", "score": 0.9},
        {"id": "b", "text": "向量命中B", "score": 0.8},
    ]
    keyword = [
        {"id": "b", "text": "关键词命中B", "score": 1.0},
        {"id": "c", "text": "关键词命中C", "score": 0.5},
    ]
    fused = rrf_fuse(vector, keyword, top_k=3)
    assert fused[0]["id"] == "b"


def test_keyword_rank_hits_chinese_bigrams():
    rows = [
        {"id": "1", "text": "主升第一天确认后才加仓", "source": "a.md"},
        {"id": "2", "text": "无关内容", "source": "b.md"},
    ]
    hits = keyword_rank("主升第一天", rows, top_k=2)
    assert hits and hits[0]["id"] == "1"


def test_kb_document_crud(isolated_kb):
    created = save_kb_document("试验规则.md", "# 试验\n\n只做纪律说明。", create=True)
    assert created["created"] is True
    docs = list_kb_documents()
    assert any(d["path"] == "试验规则.md" for d in docs)
    doc = get_kb_document("试验规则.md")
    assert "纪律说明" in doc["content"]
    preview = preview_kb_chunks("试验规则.md")
    assert preview["count"] >= 1
    save_kb_document("试验规则.md", "# 试验\n\n更新后的正文。")
    assert "更新后" in get_kb_document("试验规则.md")["content"]
    with pytest.raises(FileExistsError):
        save_kb_document("试验规则.md", "x", create=True)
    with pytest.raises(ValueError):
        save_kb_document("../escape.md", "x")
    delete_kb_document("试验规则.md")
    assert list_kb_documents() == []


def test_extract_common_formats():
    assert "只做纪律" in extract_markdown("a.txt", "只做纪律说明。".encode("utf-8"))
    html = "<html><body><h1>五灯</h1><p>红灯越多仓位越低。</p></body></html>".encode("utf-8")
    md = extract_markdown("rules.html", html)
    assert "五灯" in md and "仓位" in md
    csv_md = extract_markdown("t.csv", "灯,上限\n0,8成\n".encode("utf-8"))
    assert "8成" in csv_md
    assert md_path_for_upload("我的 研报.PDF") == "我的_研报.md"


def test_upload_txt_and_reject_unknown(isolated_kb):
    r = upload_kb_document(filename="自定义纪律.txt", data="主升第一天才能加仓。".encode("utf-8"))
    assert r["path"] == "自定义纪律.md"
    assert "主升第一天" in get_kb_document(r["path"])["content"]
    with pytest.raises(FileExistsError):
        upload_kb_document(filename="自定义纪律.txt", data="另一份".encode("utf-8"))
    again = upload_kb_document(filename="自定义纪律.txt", data="覆盖后的正文。".encode("utf-8"), overwrite=True)
    assert "覆盖后" in get_kb_document(again["path"])["content"]
    with pytest.raises(ValueError):
        upload_kb_document(filename="x.exe", data=b"MZ")


def test_upload_docx_xlsx(isolated_kb):
    from io import BytesIO

    from docx import Document
    from openpyxl import Workbook

    doc = Document()
    doc.add_heading("持仓预警", level=1)
    doc.add_paragraph("破20日线要复核。")
    buf = BytesIO()
    doc.save(buf)
    r = upload_kb_document(filename="预警.docx", data=buf.getvalue())
    text = get_kb_document(r["path"])["content"]
    assert "持仓预警" in text and "破20日线" in text

    wb = Workbook()
    ws = wb.active
    ws.title = "仓位"
    ws.append(["红灯", "上限"])
    ws.append([3, "1成"])
    xbuf = BytesIO()
    wb.save(xbuf)
    xr = upload_kb_document(filename="仓位表.xlsx", data=xbuf.getvalue())
    xt = get_kb_document(xr["path"])["content"]
    assert "1成" in xt


def test_collect_chunks_reads_markdown(isolated_kb):
    (settings.knowledge_dir / "角色.md").write_text("# 角色\n\nJarvis 是交易参谋。", encoding="utf-8")
    chunks = collect_chunks()
    assert any(c.source == "角色.md" for c in chunks)
    meta = write_kb_meta({"chunks": len(chunks)})
    assert meta["embedding"]["backend"] in ("hash", "openai")
