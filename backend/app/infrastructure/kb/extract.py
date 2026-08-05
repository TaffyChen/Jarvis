"""把常见办公/文本文件抽成 Markdown，供知识库入库。"""
from __future__ import annotations

import csv
import io
import re
from html import unescape
from html.parser import HTMLParser
from pathlib import Path

SUPPORTED_UPLOAD_EXTS = (
    ".md",
    ".markdown",
    ".txt",
    ".text",
    ".html",
    ".htm",
    ".csv",
    ".tsv",
    ".pdf",
    ".docx",
    ".xlsx",
)

_MAX_TABLE_ROWS = 200
_MAX_TABLE_COLS = 40
_MAX_PDF_PAGES = 80


def detect_upload_format(filename: str) -> str:
    ext = Path(filename or "").suffix.lower()
    if ext not in SUPPORTED_UPLOAD_EXTS:
        pretty = "、".join(SUPPORTED_UPLOAD_EXTS)
        raise ValueError(f"暂不支持 {ext or '无扩展名'}，请上传：{pretty}")
    if ext in (".md", ".markdown"):
        return "md"
    if ext in (".txt", ".text"):
        return "txt"
    if ext in (".html", ".htm"):
        return "html"
    if ext == ".tsv":
        return "tsv"
    return ext.lstrip(".")


def extract_markdown(filename: str, data: bytes) -> str:
    if not data:
        raise ValueError("空文件")
    fmt = detect_upload_format(filename)
    if fmt == "md":
        text = _decode_text(data).strip()
    elif fmt == "txt":
        text = _txt_to_md(filename, _decode_text(data))
    elif fmt == "html":
        text = _html_to_md(filename, _decode_text(data))
    elif fmt in ("csv", "tsv"):
        text = _csv_to_md(filename, _decode_text(data), dialect="excel-tab" if fmt == "tsv" else "excel")
    elif fmt == "pdf":
        text = _pdf_to_md(filename, data)
    elif fmt == "docx":
        text = _docx_to_md(filename, data)
    elif fmt == "xlsx":
        text = _xlsx_to_md(filename, data)
    else:
        raise ValueError(f"暂不支持 {fmt}")
    text = (text or "").strip()
    if not text:
        raise ValueError("未能抽出正文（扫描版 PDF / 空文档请先转成文字）")
    return text + "\n"


def md_path_for_upload(filename: str) -> str:
    stem = Path(filename or "upload").stem
    stem = re.sub(r"[^\w\u4e00-\u9fff.\-]+", "_", stem).strip("._") or "upload"
    return f"{stem}.md"


def _title_from_filename(filename: str) -> str:
    return Path(filename or "upload").stem.strip() or "upload"


def _decode_text(data: bytes) -> str:
    for enc in ("utf-8-sig", "utf-8", "gb18030", "big5"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _with_origin(filename: str, fmt: str, body: str) -> str:
    title = _title_from_filename(filename)
    head = f"<!-- jarvis-upload: original={Path(filename).name} format={fmt} -->\n# {title}\n"
    body = (body or "").strip()
    if body.startswith("#"):
        return f"<!-- jarvis-upload: original={Path(filename).name} format={fmt} -->\n{body}"
    return f"{head}\n{body}"


def _txt_to_md(filename: str, text: str) -> str:
    return _with_origin(filename, "txt", text.replace("\r\n", "\n").strip())


class _HTMLText(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []
        self._skip = 0

    def handle_starttag(self, tag: str, attrs) -> None:
        if tag in {"script", "style", "noscript"}:
            self._skip += 1
            return
        if tag in {"br", "p", "div", "tr", "li", "h1", "h2", "h3", "h4", "section"}:
            self.parts.append("\n")
        if tag in {"h1", "h2", "h3", "h4"}:
            level = int(tag[1])
            self.parts.append("#" * level + " ")

    def handle_endtag(self, tag: str) -> None:
        if tag in {"script", "style", "noscript"} and self._skip:
            self._skip -= 1
            return
        if tag in {"p", "div", "tr", "li", "h1", "h2", "h3", "h4", "section"}:
            self.parts.append("\n")

    def handle_data(self, data: str) -> None:
        if self._skip:
            return
        text = unescape(data)
        if text.strip():
            self.parts.append(text)


def _html_to_md(filename: str, html: str) -> str:
    parser = _HTMLText()
    parser.feed(html or "")
    parser.close()
    text = re.sub(r"\n{3,}", "\n\n", "".join(parser.parts)).strip()
    return _with_origin(filename, "html", text)


def _rows_to_md_table(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    width = max(len(r) for r in rows)
    norm = [(r + [""] * width)[:width] for r in rows]
    esc = [[c.replace("|", "\\|").replace("\n", " ") for c in row] for row in norm]
    header = esc[0]
    body = esc[1:] or [[""] * width]
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join("---" for _ in header) + " |",
    ]
    for row in body:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def _csv_to_md(filename: str, text: str, *, dialect: str) -> str:
    reader = csv.reader(io.StringIO(text), dialect=dialect)
    rows: list[list[str]] = []
    for i, row in enumerate(reader):
        if i >= _MAX_TABLE_ROWS:
            rows.append(["…(已截断)"])
            break
        cells = [str(c).strip() for c in row[:_MAX_TABLE_COLS]]
        if any(cells):
            rows.append(cells)
    return _with_origin(filename, "csv", _rows_to_md_table(rows))


def _pdf_to_md(filename: str, data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages: list[str] = []
    for i, page in enumerate(reader.pages[:_MAX_PDF_PAGES], start=1):
        text = (page.extract_text() or "").strip()
        if text:
            pages.append(f"## 第 {i} 页\n\n{text}")
    return _with_origin(filename, "pdf", "\n\n".join(pages))


def _docx_to_md(filename: str, data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue
        style = (para.style.name if para.style else "") or ""
        if style.startswith("Heading"):
            digits = "".join(ch for ch in style if ch.isdigit()) or "2"
            level = min(max(int(digits), 1), 6)
            parts.append("#" * level + " " + text)
        else:
            parts.append(text)
    for table in doc.tables:
        rows = []
        for row in table.rows[:_MAX_TABLE_ROWS]:
            rows.append([cell.text.strip().replace("\n", " ") for cell in row.cells[:_MAX_TABLE_COLS]])
        md = _rows_to_md_table(rows)
        if md:
            parts.append(md)
    return _with_origin(filename, "docx", "\n\n".join(parts))


def _xlsx_to_md(filename: str, data: bytes) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts: list[str] = []
    try:
        for sheet in wb.worksheets:
            rows: list[list[str]] = []
            for i, row in enumerate(sheet.iter_rows(values_only=True)):
                if i >= _MAX_TABLE_ROWS:
                    rows.append(["…(已截断)"])
                    break
                cells = [("" if c is None else str(c).replace("\n", " ").strip()) for c in row[:_MAX_TABLE_COLS]]
                if any(cells):
                    rows.append(cells)
            if not rows:
                continue
            parts.append(f"## {sheet.title}\n\n{_rows_to_md_table(rows)}")
    finally:
        wb.close()
    return _with_origin(filename, "xlsx", "\n\n".join(parts))
